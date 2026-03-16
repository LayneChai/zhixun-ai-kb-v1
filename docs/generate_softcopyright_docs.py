from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
USAGE_DOC = DOCS_DIR / "智询知识库问答系统_使用说明文档.docx"
CODE_DOC = DOCS_DIR / "智询知识库问答系统_代码(5979).docx"

SOURCE_EXTENSIONS = {".java", ".js", ".vue", ".sql", ".yml", ".yaml"}
CODE_EXTENSIONS = {".java", ".js", ".vue"}
EXCLUDE_PARTS = {".git", "node_modules", "target", "dist", "docs"}


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    footer._p.append(fld)


def add_title(doc: Document, title: str, subtitle: str, date_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_east_asia_font(run, "黑体")
    run.bold = True
    run.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(15)

    for _ in range(6):
        doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("软件名称", "智询知识库问答系统"),
        ("版本号", "V1.0"),
        ("文档名称", subtitle),
        ("适用对象", "软件著作权登记材料"),
        ("编制日期", date_text),
    ]
    for index, (left, right) in enumerate(rows):
        table.cell(index, 0).text = left
        table.cell(index, 1).text = right

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(14)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_east_asia_font(run, "黑体")
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)


def add_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(12)


def add_bullets(doc: Document, lines: Iterable[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        fmt = p.paragraph_format
        fmt.line_spacing = 1.3
        run = p.add_run(line)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12)


def add_numbered(doc: Document, lines: Iterable[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Number")
        fmt = p.paragraph_format
        fmt.line_spacing = 1.3
        run = p.add_run(line)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def get_source_files(extensions: set[str]) -> list[Path]:
    files: list[Path] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        if any(part in EXCLUDE_PARTS for part in path.parts):
            continue
        if path.suffix.lower() in extensions:
            files.append(path)
    return sorted(files, key=lambda p: str(p.relative_to(ROOT)).replace("\\", "/"))


def read_lines(path: Path) -> list[str]:
    try:
        return path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig", errors="ignore").splitlines()


def collect_code_lines() -> tuple[list[tuple[int, str]], int]:
    files = get_source_files(CODE_EXTENSIONS)
    collected: list[tuple[int, str]] = []
    line_no = 1
    for path in files:
        rel = str(path.relative_to(ROOT)).replace("\\", "/")
        collected.append((line_no, f"// ===== FILE: {rel} ====="))
        line_no += 1
        for line in read_lines(path):
            collected.append((line_no, line.replace("\t", "    ")))
            line_no += 1
    return collected, len(files)


def build_usage_doc() -> None:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])
    add_title(doc, "智询知识库问答系统", "软件使用说明文档", "2026年3月")

    add_heading(doc, "1 软件概述")
    add_paragraph(
        doc,
        "智询知识库问答系统 V1.0 是一套面向企业知识管理与智能问答场景的软件系统。"
        "系统支持知识库创建、文档上传、文本解析、知识检索、模型问答、答案来源追溯以及后台管理等功能，"
        "可用于企业内部制度问答、产品知识问答、业务资料检索、项目经验沉淀等场景。",
    )
    add_paragraph(
        doc,
        "本系统采用前后端分离架构，后端基于 Spring Boot 3、MyBatis-Plus、Sa-Token、MySQL 实现，"
        "前端基于 Vue 3、Vite、Element Plus 实现。系统围绕“上传文档、构建知识、发起问答、查看依据”"
        "的业务闭环设计，具备较强的实用性和可扩展性。",
    )
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("软件名称", "智询知识库问答系统"),
            ("版本号", "V1.0"),
            ("软件类型", "企业知识库与智能问答应用软件"),
            ("开发语言", "Java、JavaScript、Vue 单文件组件"),
            ("运行方式", "浏览器访问 + 服务端部署"),
            ("适用对象", "企业管理员、普通问答用户、系统维护人员"),
        ],
    )

    add_heading(doc, "2 运行环境")
    add_paragraph(doc, "系统运行环境建议如下。")
    add_table(
        doc,
        ["类别", "建议配置"],
        [
            ("服务器操作系统", "Windows Server 2019/2022 或 Linux 64 位发行版"),
            ("JDK", "JDK 17 及以上"),
            ("数据库", "MySQL 8.0 及以上"),
            ("缓存组件", "Redis 6.0 及以上"),
            ("Node.js", "Node.js 18 及以上"),
            ("浏览器", "Chrome、Edge、Firefox 等现代浏览器"),
        ],
    )
    add_paragraph(
        doc,
        "硬件方面，建议服务器至少具备 4 核 CPU、8GB 内存、100GB 可用磁盘空间。"
        "如需承载较大知识库文件或并发问答请求，可根据业务规模扩容数据库、缓存与应用服务器资源。",
    )

    add_heading(doc, "3 部署与启动", 1)
    add_heading(doc, "3.1 数据库初始化", 2)
    add_numbered(
        doc,
        [
            "创建数据库实例，例如数据库名为 zhixun_kb。",
            "执行 backend/src/main/resources/schema.sql 脚本，初始化系统表、知识库表、问答记录表及 RBAC 权限表。",
            "核对默认管理员账号是否已初始化成功。",
        ],
    )
    add_heading(doc, "3.2 后端服务启动", 2)
    add_numbered(
        doc,
        [
            "进入 backend 目录。",
            "根据实际环境修改 application.yml 中的数据库、Redis、端口及模型配置。",
            "执行 mvn spring-boot:run 启动服务。",
            "启动成功后，可通过 http://localhost:8080 访问接口服务。",
        ],
    )
    add_heading(doc, "3.3 前端服务启动", 2)
    add_numbered(
        doc,
        [
            "进入 frontend 目录。",
            "执行 npm install 安装依赖。",
            "执行 npm run dev 启动开发服务。",
            "浏览器访问 http://localhost:5173 进入系统界面。",
        ],
    )

    add_heading(doc, "4 总体功能说明")
    add_paragraph(
        doc,
        "系统主要由登录认证、仪表盘、知识库管理、文档管理、智能问答、大模型配置、用户管理、角色权限、菜单管理和操作日志等模块构成。"
        "各模块相互协同，形成从知识入库到问答应用再到权限审计的完整业务链路。",
    )
    add_table(
        doc,
        ["功能模块", "主要功能"],
        [
            ("登录认证", "账号密码登录、获取当前用户信息、修改个人信息、修改密码、退出登录"),
            ("知识库管理", "新建知识库、编辑知识库、启用/停用、设置默认知识库、查看知识库内容"),
            ("文档管理", "上传文档、查看文档列表、解析切片、重建索引、删除文档"),
            ("智能问答", "创建会话、历史问答、流式回答、非流式回答、来源追溯"),
            ("模型配置", "维护多套模型接入配置、启用禁用、设置默认模型"),
            ("系统管理", "用户管理、角色管理、菜单权限管理、操作日志查看、统计分析"),
        ],
    )

    add_heading(doc, "5 操作说明")
    add_heading(doc, "5.1 登录与退出", 2)
    add_paragraph(
        doc,
        "用户打开系统后进入登录页面，在账号和密码输入框中填写系统分配的登录信息，点击“立即登录”按钮即可进入系统。"
        "系统登录成功后，会根据当前用户角色跳转到相应页面。管理员用户默认进入仪表盘页面，普通用户可直接进入智能问答页面。",
    )
    add_paragraph(
        doc,
        "在已登录状态下，用户可通过页面右上角或侧边栏中的退出按钮退出系统。退出后，系统将清除当前登录令牌并返回登录页。",
    )

    add_heading(doc, "5.2 仪表盘", 2)
    add_paragraph(
        doc,
        "仪表盘用于集中展示系统运行概况。页面展示知识库总数、文档总数、问答交互总数等统计指标，同时以图表方式展示问答活跃趋势和存储资产占比。"
        "管理员可以通过仪表盘快速掌握系统运行状态，并通过快捷入口进入知识库管理、文档上传、智能问答和模型配置等页面。",
    )

    add_heading(doc, "5.3 知识库管理", 2)
    add_paragraph(
        doc,
        "知识库管理页面支持新建、编辑、启用、停用、查看内容和删除知识库。管理员点击“新增知识库”按钮后，输入知识库名称和描述信息即可完成创建。"
        "系统允许设置默认知识库，用于普通问答场景下的自动选择。",
    )
    add_paragraph(
        doc,
        "点击“查看内容”可查看指定知识库下已上传的文档及解析后生成的文本片段，便于管理员确认知识内容是否已正确入库。",
    )

    add_heading(doc, "5.4 文档管理", 2)
    add_paragraph(
        doc,
        "文档管理页面用于维护知识库中的原始资料。用户先选择目标知识库，再通过上传按钮导入 PDF、DOCX、TXT、Markdown、CSV 等文件。"
        "上传成功后，文档将记录在系统中，状态初始为“已上传”。",
    )
    add_paragraph(
        doc,
        "管理员可点击“解析切片”按钮触发文档解析。系统会自动提取文本内容，并按固定长度和重叠策略切分为多个知识片段，随后写入知识片段表。"
        "若文档内容有更新，也可点击“重建索引”以刷新检索缓存状态。",
    )

    add_heading(doc, "5.5 智能问答", 2)
    add_paragraph(
        doc,
        "智能问答页面是本系统的核心业务页面。用户可新建会话、切换历史会话、重命名会话、删除会话，并在输入框中提交自然语言问题。"
        "系统支持流式回答和普通回答两种模式。",
    )
    add_paragraph(
        doc,
        "当用户发起提问时，系统会先根据当前会话关联的知识库执行 BM25 全文检索，召回最相关的知识片段，再将问题和知识片段一并发送给大模型进行回答生成。"
        "回答结果返回后，界面会同步展示知识来源，包括文档名称、相关度评分和引用片段内容。",
    )
    add_paragraph(
        doc,
        "在流式模式下，系统采用 SSE 服务端推送方式逐段返回回答文本，用户可以实时查看生成过程；在非流式模式下，系统一次性返回完整答案。",
    )

    add_heading(doc, "5.6 大模型配置", 2)
    add_paragraph(
        doc,
        "大模型配置页面用于维护模型服务接入参数。管理员可以新增、编辑和删除模型配置项，配置内容包括供应商标识、显示名称、BaseURL、API Key、模型名称、启用状态等。"
        "系统允许维护多套配置，并可指定其中一套为默认配置，用于问答服务的实际调用。",
    )

    add_heading(doc, "5.7 用户管理", 2)
    add_paragraph(
        doc,
        "用户管理页面用于维护系统账号。管理员可查看用户列表、搜索用户、新建用户、编辑用户信息、设置用户状态并为用户分配角色。"
        "通过角色分配机制，可控制用户在系统中的菜单访问范围和操作权限。",
    )

    add_heading(doc, "5.8 角色与菜单权限管理", 2)
    add_paragraph(
        doc,
        "角色管理页面用于定义系统角色并为角色授权菜单。管理员可新建角色、修改角色名称和角色标识，并通过右侧菜单树勾选该角色可访问的菜单和功能。"
        "菜单管理页面用于维护系统导航结构及权限标识，可新建目录、菜单和按钮权限节点，从而实现细粒度授权。",
    )

    add_heading(doc, "5.9 操作日志", 2)
    add_paragraph(
        doc,
        "操作日志页面记录系统接口访问后的关键操作信息，包括模块名称、请求动作、操作者、来源 IP 和操作时间。"
        "该模块可用于系统审计、故障排查和安全追踪。",
    )

    add_heading(doc, "6 业务处理流程")
    add_numbered(
        doc,
        [
            "管理员创建知识库并设置默认知识库。",
            "管理员上传业务文档，系统保存原始文件信息。",
            "管理员执行文档解析，系统提取文本并生成知识片段。",
            "系统按知识库维度构建 BM25 检索缓存。",
            "用户创建问答会话并提交问题。",
            "系统先召回相关知识片段，再调用大模型生成答案。",
            "系统保存问答记录与来源引用信息，前端展示回答内容及来源。",
        ],
    )

    add_heading(doc, "7 数据结构说明")
    add_table(
        doc,
        ["数据表", "用途说明"],
        [
            ("sys_user", "系统用户信息"),
            ("sys_role", "系统角色信息"),
            ("sys_menu", "系统菜单与权限信息"),
            ("sys_user_role", "用户与角色关联关系"),
            ("sys_role_menu", "角色与菜单关联关系"),
            ("kb_dataset", "知识库基本信息"),
            ("kb_document", "上传文档信息"),
            ("kb_chunk", "知识片段信息"),
            ("qa_session", "问答会话信息"),
            ("qa_message", "问答消息及来源信息"),
            ("llm_config", "模型接入配置"),
            ("sys_log", "系统操作日志"),
        ],
    )
    add_paragraph(
        doc,
        "其中，kb_dataset、kb_document、kb_chunk、qa_session、qa_message 构成知识库问答业务主链路数据结构；"
        "sys_user、sys_role、sys_menu 及其关联表构成系统权限管理结构；llm_config 用于维护模型服务配置。",
    )

    add_heading(doc, "8 异常处理与注意事项")
    add_bullets(
        doc,
        [
            "若登录失败，请检查账号密码是否正确以及后端认证服务是否已启动。",
            "若上传文件失败，请检查目标知识库是否已选择、文件大小是否超过限制、上传目录是否具备写入权限。",
            "若文档解析失败，请检查文件格式是否受支持以及原始文件内容是否可正常读取。",
            "若问答无法返回结果，请检查默认知识库、默认模型配置及模型服务连通性。",
            "若前端页面出现无权限访问提示，请检查当前用户角色及对应菜单权限配置。",
        ],
    )

    add_heading(doc, "9 结论")
    add_paragraph(
        doc,
        "智询知识库问答系统 V1.0 已实现知识库管理、文档解析、BM25 检索、模型问答、答案来源追溯以及后台权限管理等完整功能。"
        "该系统具备独立运行能力，功能结构清晰，业务流程完整，可作为企业知识管理与智能问答场景的软件产品进行推广和应用。",
    )

    doc.save(USAGE_DOC)


def build_code_doc() -> None:
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    add_page_number(section)

    lines, file_count = collect_code_lines()
    if len(lines) > 3000:
        selected = lines[:1500] + lines[-1500:]
    else:
        selected = lines

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智询知识库问答系统 V1.0 源程序文档")
    set_east_asia_font(run, "黑体")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"整理说明：共统计 {file_count} 个源文件，按软著常用格式选取前 30 页与后 30 页，每页 50 行。"
    )
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(10.5)
    subtitle.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    for index, (line_no, text) in enumerate(selected, start=1):
        if index == 1501:
            marker = doc.add_paragraph()
            marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = marker.add_run("以下为源程序后 30 页")
            set_east_asia_font(run, "黑体")
            run.bold = True
            run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{line_no:05d}: {text}")
        set_east_asia_font(run, "Courier New")
        run.font.size = Pt(8.5)

        if index % 50 == 0 and index != len(selected):
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

    doc.save(CODE_DOC)


def main() -> None:
    build_usage_doc()
    build_code_doc()
    print(f"Generated: {USAGE_DOC}")
    print(f"Generated: {CODE_DOC}")


if __name__ == "__main__":
    main()
